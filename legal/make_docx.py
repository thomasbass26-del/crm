#!/usr/bin/env python3
"""Convert legal/*.md drafts to printable Word documents.
Output: ~/Desktop/Triskope Legal Drafts/NN - Title.docx"""
import os, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.expanduser("~/Desktop/Triskope Legal Drafts")
os.makedirs(OUT, exist_ok=True)

DOCS = [
    ("subscription-agreement.md",        "01 - Subscription Agreement"),
    ("acceptable-use-policy.md",         "02 - Acceptable Use Policy"),
    ("early-access-addendum.md",         "03 - Early Access Addendum"),
    ("cancellation-refund-data-policy.md","04 - Cancellation Refund and Data Policy"),
    ("dmca-policy.md",                   "05 - DMCA Policy"),
    ("site-privacy-policy-template.md",  "06 - Site Privacy Policy Template"),
    ("site-terms-of-use-template.md",    "07 - Site Terms of Use Template"),
]

def add_page_number(footer_par):
    """Insert 'Page X of Y' field codes."""
    footer_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    def fld(instr):
        r = footer_par.add_run()
        f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
        it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
        f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
        r._r.append(f1); r._r.append(it); r._r.append(f2)
    footer_par.add_run("Page ")
    fld("PAGE")
    footer_par.add_run(" of ")
    fld("NUMPAGES")

BOLD_RE = re.compile(r"\*\*(.+?)\*\*")

def add_runs(par, text, base_size=11):
    """Split **bold** markers into runs."""
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            par.add_run(text[pos:m.start()])
        b = par.add_run(m.group(1)); b.bold = True
        pos = m.end()
    if pos < len(text):
        par.add_run(text[pos:])
    for r in par.runs:
        r.font.size = Pt(base_size)

def convert(md_path, title, out_path):
    raw = open(md_path, encoding="utf-8").read()
    raw = re.sub(r"<!--.*?-->", "", raw, flags=re.S)      # strip HTML comments
    lines = raw.splitlines()

    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Inches(1))
    style = doc.styles["Normal"]
    style.font.name = "Georgia"; style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    add_page_number(sec.footer.paragraphs[0])
    fr = sec.footer.paragraphs[0].runs
    for r in fr: r.font.size = Pt(9)

    skip_h1 = True
    for ln in lines:
        s = ln.rstrip()
        if not s.strip():
            continue
        if s.startswith("# ") and skip_h1:
            skip_h1 = False
            t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = t.add_run(s[2:].strip()); r.font.size = Pt(20); r.bold = True; r.font.name = "Georgia"
            d = doc.add_paragraph(); d.alignment = WD_ALIGN_PARAGRAPH.CENTER
            dr = d.add_run("DRAFT — FOR ATTORNEY REVIEW — NOT FOR EXECUTION")
            dr.font.size = Pt(10); dr.bold = True; dr.font.color.rgb = RGBColor(0xB0, 0x20, 0x20)
            e = doc.add_paragraph(); e.alignment = WD_ALIGN_PARAGRAPH.CENTER
            er = e.add_run("Triskope LLC  ·  Myrtle Beach, South Carolina  ·  The Market Edge")
            er.font.size = Pt(9.5); er.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            continue

        if s.startswith("## "):
            h = doc.add_heading(level=1)
            r = h.add_run(s[3:].strip())
            r.font.size = Pt(13.5); r.font.name = "Georgia"; r.bold = True
            r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x22)
            continue
        if s.strip() == "---":
            doc.add_paragraph()
            continue
        if s.lstrip().startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, s.lstrip()[2:])
            continue
        # DRAFT banner line inside body: already rendered under the title
        if s.strip().startswith("**DRAFT"):
            continue
        p = doc.add_paragraph()
        add_runs(p, s.strip())

    doc.save(out_path)
    return out_path

if __name__ == "__main__":
    for fname, title in DOCS:
        out = os.path.join(OUT, f"{title}.docx")
        convert(os.path.join(SRC, fname), title, out)
        print("wrote", out)
    print("\nDone:", OUT)
